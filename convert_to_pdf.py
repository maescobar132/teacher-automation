#!/usr/bin/env python3
"""
Convert all submissions in a directory to PDF format for manual review.

Supports: .docx, .doc, .pdf (passthrough), .txt, .rtf

Usage:
    python convert_to_pdf.py --dir ~/Downloads/a --output ~/Downloads/a_pdf
    python convert_to_pdf.py --dir ~/Downloads/a.zip --strip-images

Requirements:
    - LibreOffice (for DOCX/DOC/TXT/RTF conversion)
    - python-docx (for image stripping from DOCX)
"""

from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path


SUPPORTED_EXTENSIONS = {".docx", ".doc", ".pdf", ".txt", ".rtf", ".odt"}


def strip_images_from_docx(input_path: Path, output_path: Path) -> bool:
    """
    Remove all images from a DOCX file while preserving text and formatting.

    Args:
        input_path: Path to the original DOCX
        output_path: Path to save the stripped DOCX

    Returns:
        True if successful, False otherwise
    """
    try:
        import docx
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        doc = docx.Document(input_path)

        # Remove inline images from paragraphs
        for para in doc.paragraphs:
            for run in para.runs:
                # Find and remove drawing elements (images)
                drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                for drawing in drawings:
                    drawing.getparent().remove(drawing)

                # Find and remove picture elements
                pictures = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
                for pict in pictures:
                    pict.getparent().remove(pict)

        # Remove images from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                            for drawing in drawings:
                                drawing.getparent().remove(drawing)

        doc.save(output_path)
        return True

    except Exception as e:
        print(f"    Warning: Could not strip images: {e}")
        # Copy original if stripping fails
        shutil.copy2(input_path, output_path)
        return False


def convert_doc_to_docx(input_path: Path, output_dir: Path) -> Path | None:
    """
    Convert .doc to .docx using LibreOffice.

    Returns:
        Path to converted file, or None if failed
    """
    try:
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to", "docx",
                "--outdir", str(output_dir),
                str(input_path),
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )

        if result.returncode == 0:
            docx_path = output_dir / (input_path.stem + ".docx")
            if docx_path.exists():
                return docx_path

    except Exception:
        pass

    return None


def get_output_name(input_path: Path) -> str:
    """
    Get the output filename by keeping only the part before the first underscore.

    Examples:
        John_Smith_Assignment1.docx -> John.pdf
        12345_submission_final.doc -> 12345.pdf
        NoUnderscore.docx -> NoUnderscore.pdf
    """
    stem = input_path.stem
    if "_" in stem:
        stem = stem.split("_")[0]
    return f"{stem}.pdf"


def convert_to_pdf(input_path: Path, output_dir: Path, strip_images: bool = False) -> tuple[bool, str]:
    """
    Convert a single file to PDF.

    Args:
        input_path: Path to the input file
        output_dir: Directory to save the PDF
        strip_images: If True, remove images before conversion

    Returns:
        Tuple of (success, message)
    """
    ext = input_path.suffix.lower()

    # PDF passthrough - just copy (can't easily strip images from PDF)
    if ext == ".pdf":
        output_path = output_dir / get_output_name(input_path)
        try:
            shutil.copy2(input_path, output_path)
            return True, f"Copied: {output_path.name}"
        except Exception as e:
            return False, f"Copy failed: {e}"

    # For DOCX/DOC with strip_images, process images first
    file_to_convert = input_path

    if strip_images and ext in (".docx", ".doc"):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)

            # If .doc, convert to .docx first
            if ext == ".doc":
                docx_path = convert_doc_to_docx(input_path, tmpdir_path)
                if docx_path:
                    file_to_convert = docx_path
                else:
                    # Fall back to direct conversion without stripping
                    pass
            else:
                file_to_convert = input_path

            # Strip images from DOCX
            if file_to_convert.suffix.lower() == ".docx":
                stripped_path = tmpdir_path / f"{input_path.stem}_stripped.docx"
                strip_images_from_docx(file_to_convert, stripped_path)
                file_to_convert = stripped_path

            # Convert stripped file to PDF
            try:
                result = subprocess.run(
                    [
                        "libreoffice",
                        "--headless",
                        "--convert-to", "pdf",
                        "--outdir", str(tmpdir_path),
                        str(file_to_convert),
                    ],
                    capture_output=True,
                    text=True,
                    timeout=120,
                )

                if result.returncode == 0:
                    # Find the PDF and move to output
                    pdf_files = list(tmpdir_path.glob("*.pdf"))
                    if pdf_files:
                        final_pdf = output_dir / get_output_name(input_path)
                        shutil.move(str(pdf_files[0]), str(final_pdf))
                        return True, f"Converted (images stripped): {final_pdf.name}"

            except Exception as e:
                return False, f"Error: {e}"

    # Standard conversion without image stripping
    try:
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(output_dir),
                str(input_path),
            ],
            capture_output=True,
            text=True,
            timeout=120,
        )

        if result.returncode == 0:
            # LibreOffice outputs with original stem, rename to our format
            lo_output = output_dir / f"{input_path.stem}.pdf"
            final_output = output_dir / get_output_name(input_path)
            if lo_output.exists() and lo_output != final_output:
                shutil.move(str(lo_output), str(final_output))
            return True, f"Converted: {final_output.name}"
        else:
            return False, f"LibreOffice error: {result.stderr[:200]}"

    except FileNotFoundError:
        return False, "LibreOffice not installed. Install with: sudo apt install libreoffice"
    except subprocess.TimeoutExpired:
        return False, "Conversion timed out"
    except Exception as e:
        return False, f"Error: {e}"


def convert_directory(input_dir: Path, output_dir: Path, strip_images: bool = False) -> dict:
    """
    Convert all supported files in a directory to PDF.

    Args:
        input_dir: Directory with source files
        output_dir: Directory for PDF output
        strip_images: If True, remove images from DOCX/DOC before conversion

    Returns:
        Summary dictionary with results
    """
    # Find all supported files
    files = []
    for ext in SUPPORTED_EXTENSIONS:
        files.extend(input_dir.glob(f"*{ext}"))
        files.extend(input_dir.glob(f"*{ext.upper()}"))

    files = sorted(set(files))

    if not files:
        print(f"No supported files found in {input_dir}")
        return {"total": 0, "success": 0, "failed": 0}

    # Create output directory
    output_dir.mkdir(parents=True, exist_ok=True)

    mode = " (stripping images)" if strip_images else ""
    print(f"Converting {len(files)} files to PDF{mode}...")
    print(f"Output: {output_dir}\n")

    results = {"total": len(files), "success": 0, "failed": 0, "errors": []}

    for i, file_path in enumerate(files, 1):
        print(f"[{i}/{len(files)}] {file_path.name}...", end=" ", flush=True)

        success, message = convert_to_pdf(file_path, output_dir, strip_images=strip_images)

        if success:
            results["success"] += 1
            print(f"✓ {message}")
        else:
            results["failed"] += 1
            results["errors"].append({"file": file_path.name, "error": message})
            print(f"✗ {message}")

    return results


def main():
    parser = argparse.ArgumentParser(
        description="Convert all submissions to PDF for manual review",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )

    parser.add_argument(
        "-d", "--dir",
        required=True,
        help="Directory with submission files (or ZIP file)",
    )
    parser.add_argument(
        "-o", "--output",
        help="Output directory for PDFs (default: <dir>_pdf)",
    )
    parser.add_argument(
        "--strip-images",
        action="store_true",
        help="Remove images from DOCX/DOC files before conversion (reduces file size)",
    )

    args = parser.parse_args()

    # Handle input
    input_path = Path(args.dir).expanduser().resolve()

    if not input_path.exists():
        print(f"Error: Not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    # Handle ZIP extraction
    if input_path.is_file() and input_path.suffix.lower() == ".zip":
        import zipfile

        extract_dir = input_path.parent / input_path.stem
        extract_dir.mkdir(exist_ok=True)

        print(f"Extracting ZIP: {input_path.name}")
        with zipfile.ZipFile(input_path, 'r') as zip_ref:
            zip_ref.extractall(extract_dir)
        print(f"Extracted to: {extract_dir}\n")

        input_dir = extract_dir
    else:
        input_dir = input_path

    # Determine output directory
    if args.output:
        output_dir = Path(args.output).expanduser().resolve()
    else:
        output_dir = input_dir.parent / f"{input_dir.name}_pdf"

    # Convert
    results = convert_directory(input_dir, output_dir, strip_images=args.strip_images)

    # Summary
    print("\n" + "=" * 50)
    print("SUMMARY")
    print("=" * 50)
    print(f"Total files: {results['total']}")
    print(f"Converted:   {results['success']}")
    print(f"Failed:      {results['failed']}")
    print(f"Output:      {output_dir}")

    if results["errors"]:
        print("\nFailed files:")
        for err in results["errors"]:
            print(f"  - {err['file']}: {err['error']}")

    sys.exit(0 if results["failed"] == 0 else 1)


if __name__ == "__main__":
    main()
