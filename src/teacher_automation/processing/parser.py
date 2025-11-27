"""
Document parsing and text extraction for various file formats.

Provides robust text extraction from PDF, DOCX, TXT, and source code files
with proper error handling, encoding detection, and metadata extraction.
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from ..utils.logging import get_logger
from .filetypes import detect_filetype, FileCategory, PDF, DOCX, DOC, TXT

logger = get_logger(__name__)


# -----------------------------------------------------------------------------
# Exceptions
# -----------------------------------------------------------------------------


class ParseError(Exception):
    """Error during document parsing."""

    pass


class UnsupportedFormatError(ParseError):
    """File format not supported for parsing."""

    pass


class EncodingError(ParseError):
    """Could not determine or decode file encoding."""

    pass


# -----------------------------------------------------------------------------
# Data Classes
# -----------------------------------------------------------------------------


@dataclass
class TextExtractionResult:
    """Result of text extraction from a document."""

    text: str
    file_path: Path
    format: str
    word_count: int = 0
    char_count: int = 0
    page_count: int = 0
    metadata: dict[str, Any] = field(default_factory=dict)
    warnings: list[str] = field(default_factory=list)

    def __post_init__(self):
        """Calculate word and character counts."""
        if not self.word_count:
            self.word_count = len(self.text.split())
        if not self.char_count:
            self.char_count = len(self.text)


@dataclass
class PDFMetadata:
    """Metadata extracted from a PDF document."""

    title: str = ""
    author: str = ""
    subject: str = ""
    creator: str = ""
    producer: str = ""
    creation_date: str = ""
    modification_date: str = ""
    page_count: int = 0


@dataclass
class DOCXMetadata:
    """Metadata extracted from a DOCX document."""

    title: str = ""
    author: str = ""
    subject: str = ""
    keywords: str = ""
    comments: str = ""
    created: str = ""
    modified: str = ""
    last_modified_by: str = ""
    revision: int = 0
    paragraph_count: int = 0
    word_count: int = 0


# -----------------------------------------------------------------------------
# Text Extraction Functions
# -----------------------------------------------------------------------------


def extract_text(file_path: Path, encoding: str | None = None) -> TextExtractionResult:
    """
    Extract text from a document file.

    Automatically detects file type and uses the appropriate extraction method.
    Supports PDF, DOCX, TXT, and source code files.

    Args:
        file_path: Path to the document
        encoding: Encoding hint for text files (auto-detected if not specified)

    Returns:
        TextExtractionResult with extracted text and metadata

    Raises:
        ParseError: If parsing fails
        UnsupportedFormatError: If file format not supported
    """
    if not file_path.exists():
        raise ParseError(f"File not found: {file_path}")

    file_type = detect_filetype(file_path)

    if file_type == PDF:
        return extract_text_from_pdf(file_path)
    elif file_type == DOCX:
        return extract_text_from_docx(file_path)
    elif file_type == DOC:
        return extract_text_from_doc(file_path)
    elif file_type.category in (FileCategory.DOCUMENT, FileCategory.CODE):
        return extract_text_from_text_file(file_path, encoding)
    else:
        raise UnsupportedFormatError(f"Unsupported file format: {file_type.extension}")


def extract_text_from_pdf(file_path: Path) -> TextExtractionResult:
    """
    Extract text from a PDF document.

    Uses pypdf for extraction. Falls back to pdfminer.six if available
    for better text extraction from complex PDFs.

    Args:
        file_path: Path to the PDF file

    Returns:
        TextExtractionResult with extracted text and PDF metadata

    Raises:
        ParseError: If PDF parsing fails
    """
    logger.info(f"Extracting text from PDF: {file_path}")
    warnings = []

    # Try pypdf first (most common)
    try:
        import pypdf

        reader = pypdf.PdfReader(file_path)

        # Extract metadata
        meta = reader.metadata or {}
        pdf_meta = PDFMetadata(
            title=str(meta.get("/Title", "")),
            author=str(meta.get("/Author", "")),
            subject=str(meta.get("/Subject", "")),
            creator=str(meta.get("/Creator", "")),
            producer=str(meta.get("/Producer", "")),
            creation_date=str(meta.get("/CreationDate", "")),
            modification_date=str(meta.get("/ModDate", "")),
            page_count=len(reader.pages),
        )

        # Extract text from all pages
        text_parts = []
        for i, page in enumerate(reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text.strip())
            except Exception as e:
                warnings.append(f"Could not extract text from page {i}: {e}")

        if not text_parts:
            warnings.append("No text could be extracted from PDF (may be scanned/image-based)")

        text = "\n\n".join(text_parts)

        return TextExtractionResult(
            text=text,
            file_path=file_path,
            format="pdf",
            page_count=pdf_meta.page_count,
            metadata={
                "title": pdf_meta.title,
                "author": pdf_meta.author,
                "subject": pdf_meta.subject,
                "creator": pdf_meta.creator,
                "producer": pdf_meta.producer,
            },
            warnings=warnings,
        )

    except ImportError:
        pass  # Try alternative

    # Try pdfminer.six as fallback
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        from pdfminer.pdfpage import PDFPage

        text = pdfminer_extract(str(file_path))

        # Count pages
        with open(file_path, "rb") as f:
            page_count = sum(1 for _ in PDFPage.get_pages(f))

        return TextExtractionResult(
            text=text.strip(),
            file_path=file_path,
            format="pdf",
            page_count=page_count,
            metadata={},
            warnings=warnings,
        )

    except ImportError:
        raise ParseError(
            "No PDF library installed. Install with: pip install pypdf "
            "or pip install pdfminer.six"
        )
    except Exception as e:
        raise ParseError(f"Failed to parse PDF: {e}") from e


def extract_text_from_doc(file_path: Path) -> TextExtractionResult:
    """
    Extract text from a legacy .doc (Word 97-2003) document.

    Converts to DOCX using LibreOffice, then extracts text.
    Falls back to antiword or catdoc if available.

    Args:
        file_path: Path to the .doc file

    Returns:
        TextExtractionResult with extracted text

    Raises:
        ParseError: If conversion or parsing fails
    """
    logger.info(f"Extracting text from DOC: {file_path}")
    import subprocess
    import tempfile

    # Try LibreOffice conversion first (most reliable)
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            # Convert to DOCX using LibreOffice
            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", tmpdir,
                    str(file_path),
                ],
                capture_output=True,
                text=True,
                timeout=60,
            )

            if result.returncode == 0:
                # Find the converted file
                docx_path = Path(tmpdir) / (file_path.stem + ".docx")
                if docx_path.exists():
                    # Extract text from the converted DOCX
                    docx_result = extract_text_from_docx(docx_path)
                    # Update the file_path to reflect original
                    return TextExtractionResult(
                        text=docx_result.text,
                        file_path=file_path,
                        format="doc",
                        word_count=docx_result.word_count,
                        char_count=docx_result.char_count,
                        metadata=docx_result.metadata,
                        warnings=["Converted from DOC to DOCX using LibreOffice"],
                    )

    except FileNotFoundError:
        logger.debug("LibreOffice not found, trying alternatives")
    except subprocess.TimeoutExpired:
        logger.warning("LibreOffice conversion timed out")
    except Exception as e:
        logger.warning(f"LibreOffice conversion failed: {e}")

    # Try antiword as fallback
    try:
        result = subprocess.run(
            ["antiword", str(file_path)],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            return TextExtractionResult(
                text=result.stdout.strip(),
                file_path=file_path,
                format="doc",
                metadata={},
                warnings=["Extracted using antiword"],
            )
    except FileNotFoundError:
        logger.debug("antiword not found")
    except Exception as e:
        logger.warning(f"antiword failed: {e}")

    # Try catdoc as last resort
    try:
        result = subprocess.run(
            ["catdoc", str(file_path)],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            return TextExtractionResult(
                text=result.stdout.strip(),
                file_path=file_path,
                format="doc",
                metadata={},
                warnings=["Extracted using catdoc"],
            )
    except FileNotFoundError:
        logger.debug("catdoc not found")
    except Exception as e:
        logger.warning(f"catdoc failed: {e}")

    raise ParseError(
        f"Cannot extract text from .doc file: {file_path}. "
        "Install LibreOffice (recommended), antiword, or catdoc."
    )


def extract_text_from_docx(file_path: Path) -> TextExtractionResult:
    """
    Extract text from a DOCX document.

    Extracts text from paragraphs, tables, headers, and footers.
    Also extracts document metadata.

    Args:
        file_path: Path to the DOCX file

    Returns:
        TextExtractionResult with extracted text and DOCX metadata

    Raises:
        ParseError: If DOCX parsing fails
    """
    logger.info(f"Extracting text from DOCX: {file_path}")
    warnings = []

    try:
        import docx
        from docx.opc.exceptions import PackageNotFoundError

        try:
            doc = docx.Document(file_path)
        except PackageNotFoundError as e:
            raise ParseError(f"Invalid or corrupted DOCX file: {file_path}") from e

        # Extract metadata
        core_props = doc.core_properties
        docx_meta = DOCXMetadata(
            title=core_props.title or "",
            author=core_props.author or "",
            subject=core_props.subject or "",
            keywords=core_props.keywords or "",
            comments=core_props.comments or "",
            created=str(core_props.created) if core_props.created else "",
            modified=str(core_props.modified) if core_props.modified else "",
            last_modified_by=core_props.last_modified_by or "",
            revision=core_props.revision or 0,
        )

        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = _extract_table_text(table)
            if table_text:
                text_parts.append(table_text)

        # Extract headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        text_parts.insert(0, f"[Header] {para.text}")

            # Footer
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        text_parts.append(f"[Footer] {para.text}")

        text = "\n\n".join(text_parts)

        return TextExtractionResult(
            text=text,
            file_path=file_path,
            format="docx",
            metadata={
                "title": docx_meta.title,
                "author": docx_meta.author,
                "subject": docx_meta.subject,
                "keywords": docx_meta.keywords,
                "last_modified_by": docx_meta.last_modified_by,
            },
            warnings=warnings,
        )

    except ImportError:
        raise ParseError("python-docx not installed. Install with: pip install python-docx")
    except Exception as e:
        raise ParseError(f"Failed to parse DOCX: {e}") from e


def _extract_table_text(table) -> str:
    """Extract text from a DOCX table."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows) if rows else ""


def extract_text_from_text_file(
    file_path: Path,
    encoding: str | None = None,
) -> TextExtractionResult:
    """
    Extract text from a plain text file with encoding detection.

    Attempts to detect encoding automatically if not specified.
    Tries UTF-8, Latin-1, and other common encodings.

    Args:
        file_path: Path to the text file
        encoding: Encoding to use (auto-detected if None)

    Returns:
        TextExtractionResult with file contents

    Raises:
        EncodingError: If encoding cannot be determined
        ParseError: If file cannot be read
    """
    logger.info(f"Extracting text from file: {file_path}")

    if encoding:
        try:
            text = file_path.read_text(encoding=encoding)
            return TextExtractionResult(
                text=text,
                file_path=file_path,
                format=file_path.suffix.lstrip(".") or "txt",
                metadata={"encoding": encoding},
            )
        except UnicodeDecodeError as e:
            raise EncodingError(f"Cannot decode file with encoding {encoding}: {e}") from e

    # Try to detect encoding
    detected_encoding = _detect_encoding(file_path)

    # List of encodings to try
    encodings_to_try = [detected_encoding] if detected_encoding else []
    encodings_to_try.extend(["utf-8", "utf-8-sig", "latin-1", "cp1252", "iso-8859-1"])

    # Remove duplicates while preserving order
    seen = set()
    encodings_to_try = [e for e in encodings_to_try if e and e not in seen and not seen.add(e)]

    for enc in encodings_to_try:
        try:
            text = file_path.read_text(encoding=enc)

            return TextExtractionResult(
                text=text,
                file_path=file_path,
                format=file_path.suffix.lstrip(".") or "txt",
                metadata={"encoding": enc, "detected": enc == detected_encoding},
            )
        except (UnicodeDecodeError, LookupError):
            continue

    # Last resort: read with errors replaced
    text = file_path.read_text(encoding="utf-8", errors="replace")

    return TextExtractionResult(
        text=text,
        file_path=file_path,
        format=file_path.suffix.lstrip(".") or "txt",
        metadata={"encoding": "utf-8", "had_errors": True},
        warnings=["Some characters could not be decoded and were replaced"],
    )


def _detect_encoding(file_path: Path) -> str | None:
    """
    Detect file encoding using chardet if available.

    Args:
        file_path: Path to the file

    Returns:
        Detected encoding or None
    """
    try:
        import chardet

        with open(file_path, "rb") as f:
            raw_data = f.read(10000)  # Read first 10KB

        result = chardet.detect(raw_data)
        if result and result.get("confidence", 0) > 0.7:
            return result.get("encoding")

    except ImportError:
        pass  # chardet not available

    # Check for BOM
    try:
        with open(file_path, "rb") as f:
            bom = f.read(4)

        if bom.startswith(b"\xef\xbb\xbf"):
            return "utf-8-sig"
        elif bom.startswith(b"\xff\xfe\x00\x00"):
            return "utf-32-le"
        elif bom.startswith(b"\x00\x00\xfe\xff"):
            return "utf-32-be"
        elif bom.startswith(b"\xff\xfe"):
            return "utf-16-le"
        elif bom.startswith(b"\xfe\xff"):
            return "utf-16-be"

    except IOError:
        pass

    return None


# -----------------------------------------------------------------------------
# Text Processing Utilities
# -----------------------------------------------------------------------------


def normalize_text(text: str) -> str:
    """
    Normalize text for consistent processing.

    - Normalizes whitespace
    - Removes excessive blank lines
    - Normalizes line endings

    Args:
        text: Raw text to normalize

    Returns:
        Normalized text
    """
    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Normalize whitespace within lines
    lines = []
    for line in text.split("\n"):
        # Collapse multiple spaces but preserve indentation
        match = re.match(r"^(\s*)", line)
        indent = match.group(1) if match else ""
        content = " ".join(line.split())
        lines.append(indent + content if content else "")

    # Remove excessive blank lines (more than 2 in a row)
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def count_words(text: str) -> int:
    """
    Count words in text.

    Args:
        text: Text to count words in

    Returns:
        Word count
    """
    # Split on whitespace and filter empty strings
    words = [w for w in text.split() if w]
    return len(words)


def extract_sentences(text: str) -> list[str]:
    """
    Extract sentences from text.

    Args:
        text: Text to extract sentences from

    Returns:
        List of sentences
    """
    # Simple sentence splitting on common terminators
    # For production, use nltk or spacy
    sentences = re.split(r"(?<=[.!?])\s+", text)
    return [s.strip() for s in sentences if s.strip()]


# -----------------------------------------------------------------------------
# Document Parser Class (Backward Compatibility)
# -----------------------------------------------------------------------------


class DocumentParser:
    """
    Parses various document formats to extract text content.

    This class provides backward compatibility with the original API.
    For new code, prefer using the extract_text() function directly.
    """

    def parse(self, file_path: Path) -> str:
        """Parse a document and return its text content.

        Args:
            file_path: Path to the document

        Returns:
            Extracted text content

        Raises:
            ValueError: If the file format is not supported
        """
        try:
            result = extract_text(file_path)
            return result.text
        except UnsupportedFormatError as e:
            raise ValueError(str(e)) from e
        except ParseError as e:
            logger.error(f"Parse error for {file_path}: {e}")
            raise

    def parse_with_metadata(self, file_path: Path) -> TextExtractionResult:
        """
        Parse a document and return text with metadata.

        Args:
            file_path: Path to the document

        Returns:
            TextExtractionResult with text and metadata
        """
        return extract_text(file_path)

    # Legacy method names for backward compatibility
    def _parse_text(self, file_path: Path) -> str:
        """Parse a plain text file."""
        return extract_text_from_text_file(file_path).text

    def _parse_code(self, file_path: Path) -> str:
        """Parse a source code file."""
        return extract_text_from_text_file(file_path).text

    def _parse_pdf(self, file_path: Path) -> str:
        """Parse a PDF file."""
        return extract_text_from_pdf(file_path).text

    def _parse_docx(self, file_path: Path) -> str:
        """Parse a Word document."""
        return extract_text_from_docx(file_path).text
